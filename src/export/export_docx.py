from docx import Document
import io

def export_docx(data):
    doc = Document()
    doc.add_heading('Báo cáo thẩm định', level=1)
    idf = data.get('identification',{})
    doc.add_paragraph(f"Họ tên: {idf.get('ten','')}")
    doc.add_paragraph(f"CCCD: {idf.get('cccd','')}")
    doc.add_paragraph('\nPhương án:')
    fin = data.get('finance',{})
    doc.add_paragraph(f"Mục đích: {fin.get('muc_dich','')}")
    doc.add_paragraph(f"Số tiền vay: {fin.get('so_tien_vay','')}")
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b.getvalue()
